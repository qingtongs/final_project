"""
Generate Chinese-only report based on 报告模板.docx.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "项目要求" / "报告模板.docx"
REPORT_DIR = PROJECT_ROOT / "项目要求" / "报告"
REPORT_PATH = REPORT_DIR / "animal_detection_project_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_from_paragraph(doc: Document, start_index: int) -> None:
    body = doc._body._element
    for paragraph in list(doc.paragraphs[start_index:]):
        body.remove(paragraph._element)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "2F5597"),
    ]:
        style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    if "List Bullet" not in [s.name for s in doc.styles]:
        bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.font.name = "Microsoft YaHei"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        bullet.font.size = Pt(10.5)


def add_paragraph(doc: Document, text: str = "", style: str | None = None):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(item)
        for run in p.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")
        header_cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = Cm(widths[i])


def add_figure(doc: Document, image_path: Path, caption: str,
               width: float = 5.8) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[图片未找到: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in c.runs:
        run.italic = True
        run.font.size = Pt(9)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def configure_cover(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Deep Learning Group Project":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.name = "Microsoft YaHei"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        elif text.startswith("(2025-2026"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("提交日期"):
            paragraph.text = "提交日期：2026年 6月 日"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert project title
    for i, p in enumerate(doc.paragraphs):
        if "提交日期" in p.text:
            body = doc._body._element
            new_p = OxmlElement("w:p")
            pPr = OxmlElement("w:pPr")
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "center")
            pPr.append(jc)
            new_p.append(pPr)
            r = OxmlElement("w:r")
            rPr = OxmlElement("w:rPr")
            rFonts = OxmlElement("w:rFonts")
            rFonts.set(qn("w:ascii"), "Microsoft YaHei")
            rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            rPr.append(rFonts)
            sz = OxmlElement("w:sz")
            sz.set(qn("w:val"), "32")
            rPr.append(sz)
            b = OxmlElement("w:b")
            rPr.append(b)
            color = OxmlElement("w:color")
            color.set(qn("w:val"), "1F4E79")
            rPr.append(color)
            r.append(rPr)
            t = OxmlElement("w:t")
            t.text = "动物检测与计数系统"
            t.set(qn("xml:space"), "preserve")
            r.append(t)
            new_p.append(r)
            p._element.addnext(new_p)
            break

    if doc.tables:
        table = doc.tables[0]
        values = [
            ["姓   名", "待填写", "学    号", "待填写"],
            ["姓   名", "待填写", "学    号", "待填写"],
            ["姓   名", "待填写", "学    号", "待填写"],
            ["专业班级", "待填写", "课程代码", "待填写"],
            ["课程名称", "Deep Learning", "任课教师", "曾子倩"],
        ]
        for row, values_row in zip(table.rows, values):
            for cell, value in zip(row.cells, values_row):
                set_cell_text(cell, value,
                             bold=value in {"姓   名", "学    号", "专业班级", "课程代码", "课程名称", "任课教师"})
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def main() -> None:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    set_document_defaults(doc)
    configure_cover(doc)

    try:
        start = next(i for i, p in enumerate(doc.paragraphs)
                     if p.text.strip().startswith("Introduction"))
    except StopIteration:
        start = len(doc.paragraphs)
    remove_from_paragraph(doc, start)

    # ============================
    # 1. 引言
    # ============================
    doc.add_page_break()
    add_paragraph(doc, "一、引言", "Heading 1")
    add_paragraph(doc,
        "本项目为深度学习期末项目实现了动物检测与计数系统。系统以单张图片作为输入，返回一个"
        "Python风格的字典，键为支持的动物类别，值为正整数计数。最终实现遵循要求的20类词汇表，"
        "并将批量预测结果写入JSON文件，不含任何自然语言解释。"
    )
    add_paragraph(doc,
        "本项目基于YOLO11s和Ultralytics训练/推理框架。选择YOLO是因为它在单次前向传播中同时"
        "完成目标定位、分类和置信度估计，非常适合有时间限制的验证和最终评估流程。项目覆盖了"
        "完整流程：数据收集与标注、模型训练与微调、检测推理、后处理和JSON预测导出。"
    )

    # ============================
    # 2. 方法
    # ============================
    add_paragraph(doc, "二、方法", "Heading 1")

    add_paragraph(doc, "2.1 数据集构建", "Heading 2")
    add_paragraph(doc,
        "训练数据通过多个阶段逐步构建。首先，合并了两个源数据集（animal_dataset_xfz和"
        "animal_dataset_yzy），并将其标注转换为统一的JSON格式。基于合并后的数据集，准备了"
        "animal_dataset1作为20类YOLO格式数据集，包含训练集和验证集划分。"
    )
    add_paragraph(doc,
        "为改善弱类别表现，构建了两个额外的定向数据集：animal_dataset_5sp_v1（聚焦狼、鹅、"
        "马、牛和猴子）和animal_dataset_6sp_v1（继续优化困难类别）。两个数据集均以LabelMe风格"
        "JSON格式自动标注，然后转换为YOLO格式，同时保留完整的20类输出头。该方法确保在定向类别"
        "上微调不会降低模型识别原始20个动物类别的能力。"
    )
    add_bullets(doc, [
        "支持类别：猫、狗、马、牛、羊、山羊、猪、兔、鸡、鸭、鹅、鹿、猴、狐狸、狼、熊、老虎、"
        "狮子、斑马、长颈鹿（共20类）。",
        "数据来源包括人工标注、带可视化预览的自动标注JSON，以及针对弱类别的定向微调数据集。",
        "最终数据集路径：dataset/animal_dataset_6sp_yolo_20cls，由configs/animal_dataset_6sp_20cls.yaml配置。",
    ])

    add_paragraph(doc, "2.2 模型训练与微调", "Heading 2")
    add_paragraph(doc,
        "基础检测器为YOLO11s（YOLO11的小型变体），使用COCO预训练权重。训练分三个阶段进行，"
        "每个阶段基于前一阶段的最佳模型："
    )
    add_table(doc,
        ["阶段", "输入数据", "目的", "输出模型"],
        [
            ["阶段1", "animal_dataset1", "从COCO基线训练20个动物类别", "animal_dataset1 best.pt"],
            ["阶段2", "animal_dataset_5sp_v1", "改善狼/鹅/马/牛/猴的识别", "5sp 20类 best.pt"],
            ["阶段3", "animal_dataset_6sp_v1", "从最佳模型继续优化", "6sp 20类 best.pt"],
        ],
        [1.5, 3.2, 5.8, 3.7],
    )
    add_paragraph(doc,
        "训练超参数：图像尺寸640×640，批次大小8，冻结前10层，早停耐心值15轮，AdamW优化器"
        "（初始学习率0.0005，权重衰减0.0005，预热3轮）。数据增强包括HSV抖动（h=0.015, "
        "s=0.7, v=0.4）、旋转（±8°）、平移（±10%）、缩放（±40%）、水平翻转（50%）、"
        "垂直翻转（5%）、马赛克（100%）和轻量mixup（5%）。马赛克增强在第10轮后关闭。"
        "训练在单张NVIDIA GeForce RTX 3070 Ti笔记本GPU上使用CUDA进行。"
    )

    add_paragraph(doc, "2.3 检测与后处理流程", "Heading 2")
    add_paragraph(doc, "推理流程按以下步骤处理每张图片：")
    add_bullets(doc, [
        "1) YOLO预测：运行微调后的YOLO11s模型，获取每个检测的边界框、置信度分数和预测类别标签。",
        "2) 类别白名单过滤：将YOLO输出标签映射到20个支持的动物类别。不支持的标签（如人、鸟、大象）被丢弃。",
        "3) 置信度阈值与NMS：丢弃置信度低于0.25的检测。YOLO内置NMS以IoU阈值0.45应用。",
        "4) 自定义重复抑制：YOLO NMS之外的额外后处理层。同类别框在IoU≥0.40时合并，或中心距离近"
        "（≤最小对角线的30%）且尺寸相似（面积比≥0.50）时合并。跨类别框在更严格条件下合并："
        "IoU≥0.62或中心距离≤最小对角线的18%且面积比≥0.70。所有情况下保留置信度最高的检测。",
        "5) 计数：按映射后的类别名称分组剩余检测，统计每个类别的实例数。",
        "6) JSON导出：为每张图片写入一个字典，仅包含具有正计数的类别。输出格式："
        "{\"eval_000001\": {\"cat\": 2, \"duck\": 1, \"deer\": 1}}。",
    ])
    add_paragraph(doc,
        "相同的重复抑制逻辑在JSON预测计数和标注图像可视化之间共享，确保显示框与提交计数之间的一致性。"
    )

    # ============================
    # 3. 评估结果
    # ============================
    add_paragraph(doc, "三、评估结果", "Heading 1")

    add_paragraph(doc, "3.1 最终模型性能", "Heading 2")
    add_paragraph(doc,
        "最终模型在第三阶段微调后，所有训练指标均表现强劲。基于验证集mAP50，第2轮被选为最佳检查点。"
    )
    add_table(doc,
        ["指标", "数值"],
        [
            ["精确率 (Precision)", "0.963"],
            ["召回率 (Recall)", "0.977"],
            ["mAP50", "0.978"],
            ["mAP50-95", "0.940"],
            ["最佳轮次", "第2轮"],
            ["模型路径", "runs/detect/.../animal_dataset_6sp_20cls_finetune/weights/best.pt"],
        ],
        [4.5, 9.5],
    )

    add_paragraph(doc, "3.2 验证集结果", "Heading 2")
    add_paragraph(doc,
        "最终内部验证使用15张图片的验证集（val_set）。使用平衡式抑制后处理配置，模型在项目"
        "评分规则下取得了平均90.94/100的评分。15张图片中有9张获得满分（100/100），表明系统"
        "正确处理了大多数检测场景。"
    )
    add_table(doc,
        ["图片组", "代表性结果"],
        [
            ["完美预测", "eval_001384, eval_001388, eval_001431, val_4, val_5, val_6, val_7, val_8, val_10"],
            ["主要漏检", "val_1漏检狐狸；eval_001391将牛的数量预测为2而非3"],
            ["主要混淆", "val_2和val_9将鸭与鹅混淆；eval_001402产生多余的马检测"],
        ],
        [3.5, 10.5],
    )

    curves_path = PROJECT_ROOT / "runs" / "detect" / "runs" / "detect" / "animal_dataset_6sp_20cls_finetune" / "results.png"
    add_figure(doc, curves_path, "图1. 最终6sp 20类微调阶段的训练和验证曲线。", 5.5)

    anno_path = PROJECT_ROOT / "outputs" / "val_set_annotated_6sp_20cls_finetuned_suppressed_balanced" / "annotated" / "eval_001384.png"
    add_figure(doc, anno_path, "图2. 带类别标签、置信度分数和边界框的验证输出示例。", 4.5)

    # ============================
    # 4. 分析
    # ============================
    add_paragraph(doc, "四、分析", "Heading 1")

    add_paragraph(doc, "4.1 性能分析", "Heading 2")
    add_paragraph(doc,
        "模型在大型、视觉特征明显的动物（如斑马、长颈鹿、马、狮子、老虎和熊）上表现优异，"
        "mAP超过0.95。这些类别受益于独特的形状、纹理和颜色模式，YOLO能够从训练数据中很好地学习。"
    )
    add_paragraph(doc,
        "剩余错误主要来自三个方面：（1）视觉相似物种，特别是鸭和鹅在合成场景中体型和颜色相似，"
        "以及狐狸、狼和狗在纹理和姿态上重叠；（2）遮挡和密集排列动物，部分可见性导致漏检或"
        "合并检测；（3）合成场景伪影可能无法完美代表真实世界的动物外观。"
    )

    add_paragraph(doc, "4.2 错误分析", "Heading 2")
    add_bullets(doc, [
        "鸭/鹅混淆：这两种禽类是最常混淆的组合。在合成图像中，两者呈现相似的白色/棕色体色和"
        "相近大小，没有更多定向训练样本的情况下，细粒度区分非常困难。",
        "犬科混淆：狐狸、狼和狗具有相似的身体结构、皮毛纹理和姿态。在远距离或非特征性姿态下，"
        "模型可能在这三个类别之间错误分类。",
        "计数错误：遮挡和密集排列仍然具有挑战性。当动物部分重叠时，模型可能将两只动物检测为"
        "一只，或当身体部位被遮挡物视觉分离时将一只动物检测为两只。",
        "误报惩罚：评分规则对每个额外错误类别扣除5分，使误报类别代价特别高。需要保守的置信度"
        "阈值来平衡召回率和引入虚假类别的风险。",
    ])

    add_paragraph(doc, "4.3 改进策略", "Heading 2")
    add_bullets(doc, [
        "专门为鸭/鹅和狐狸/狼/狗组合增加更多人工验证的训练样本，并进行仔细的标注审查以确保"
        "标注一致性。",
        "在推理过程中探索测试时增强（TTA）以提高鲁棒性，特别是针对有遮挡或不寻常视角的图像。",
        "研究逐类置信度阈值调优：对经常混淆的类别使用更高阈值，对可靠检测的类别使用更低阈值。",
        "考虑集成方法，结合多个YOLO变体或引入额外的特征提取器进行细粒度物种区分。",
    ])

    # ============================
    # 5. 贡献
    # ============================
    add_paragraph(doc, "五、贡献", "Heading 1")
    add_paragraph(doc, "下表描述了每位团队成员对项目的贡献。")
    add_table(doc,
        ["成员", "贡献"],
        [
            ["成员1\n待填写",
             "数据集收集、合并和LabelMe JSON标注格式转换。验证数据组织和预处理脚本。"],
            ["成员2\n待填写",
             "YOLO模型训练流程、CUDA环境配置、多阶段微调、超参数调优和模型性能跟踪。"],
            ["成员3\n待填写",
             "评估和推理脚本、自定义重复框后处理算法、标注图像生成、报告撰写、PPT制作和错误分析。"],
        ],
        [3.0, 11.0],
    )

    # ============================
    # 6. 结论
    # ============================
    add_paragraph(doc, "六、结论", "Heading 1")
    add_paragraph(doc,
        "本项目成功完成了一个可复现的基于YOLO的动物检测与计数流程，覆盖了完整工作流："
        "数据准备与标注、多阶段模型微调、带自定义后处理的批量推理、JSON预测导出、验证评分"
        "以及带边界框和置信度分数的可视化标注。最终20类模型取得了90.94/100的高验证评分和"
        "强劲的检测指标（mAP50: 0.978, mAP50-95: 0.940）。"
    )
    add_paragraph(doc,
        "关键技术贡献包括：（1）渐进式数据构建策略，针对弱类别同时保留完整20类词汇表；"
        "（2）平衡式重复抑制算法，减少误报但不过度删除真实邻近动物；（3）可视化标注和"
        "JSON计数之间共享后处理逻辑，确保显示结果与提交结果一致。"
    )
    add_paragraph(doc,
        "未来工作应重点关注为最易混淆的类别对（鸭/鹅、狐狸/狼/狗）收集更多人工验证的困难"
        "样本，探索逐类阈值优化，并研究测试时增强和集成方法等先进技术，以进一步提高在涉及"
        "遮挡和细粒度物种区分的挑战性案例上的鲁棒性。"
    )

    # ============================
    # 保存
    # ============================
    doc.save(str(REPORT_PATH))
    print(f"报告已保存: {REPORT_PATH}")


if __name__ == "__main__":
    main()
