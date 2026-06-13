from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
SOURCE_DIR = Path(r"C:\Users\20344\Desktop\ex1\Final Project\Final Project")
TEMPLATE_PATH = SOURCE_DIR / "报告模板.docx"
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "final_deliverables"
REPORT_PATH = OUTPUT_DIR / "animal_detection_project_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_from_paragraph(doc: Document, start_index: int) -> None:
    body = doc._body._element
    for paragraph in list(doc.paragraphs[start_index:]):
        body.remove(paragraph._element)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "2F5597"),
    ]:
        style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    if "List Bullet" not in [s.name for s in doc.styles]:
        bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.font.name = "Arial"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        bullet.font.size = Pt(10.5)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(item)
        for run in p.runs:
            run.font.name = "Arial"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")
        header_cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = Cm(widths[i])


def add_figure(doc: Document, image_path: Path, caption: str, width: float = 5.8) -> None:
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    c.runs[0].italic = True
    c.runs[0].font.size = Pt(9)


def configure_cover(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Deep Learning Group Project":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.runs[0].font.size = Pt(20)
            paragraph.runs[0].bold = True
            paragraph.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        elif text.startswith("提交日期"):
            paragraph.text = "提交日期：2026年 6月 日"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if doc.tables:
        table = doc.tables[0]
        values = [
            ["姓   名", "待填写", "学    号", "待填写"],
            ["姓   名", "待填写", "学    号", "待填写"],
            ["姓   名", "待填写", "学    号", "待填写"],
            ["专业班级", "待填写", "课程代码", "待填写"],
            ["课程名称", "Deep Learning", "任课教师", "曾子倩"],
        ]
        for row, values_row in zip(table.rows, values):
            for cell, value in zip(row.cells, values_row):
                set_cell_text(cell, value, bold=value in {"姓   名", "学    号", "专业班级", "课程代码", "课程名称", "任课教师"})
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    set_document_defaults(doc)
    configure_cover(doc)

    start = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip().startswith("Introduction"))
    remove_from_paragraph(doc, start)

    doc.add_page_break()
    add_paragraph(doc, "Introduction", "Heading 1")
    add_paragraph(
        doc,
        "This project implements an animal detection and counting system for the Deep Learning final project. "
        "The system accepts a single image and returns a clean dictionary whose keys are animal categories and whose values are positive integer counts. "
        "The final implementation follows the required 20-category vocabulary and writes batch predictions to JSON without natural-language explanations."
    )
    add_paragraph(
        doc,
        "The work is based on YOLO11s and the Ultralytics training/inference interface. YOLO was selected because it performs object localization, classification, and confidence estimation in one forward pass, which is suitable for the time-limited validation and final evaluation workflow."
    )

    add_paragraph(doc, "Methodology", "Heading 1")
    add_paragraph(doc, "Dataset construction", "Heading 2")
    add_paragraph(
        doc,
        "The training data was built progressively. First, animal_dataset_xfz and animal_dataset_yzy were merged and converted into JSON annotations. "
        "Then animal_dataset1 was prepared as a 20-class YOLO dataset. Two additional datasets, animal_dataset_5sp_v1 and animal_dataset_6sp_v1, were auto-labeled in LabelMe-style JSON format and converted into YOLO format while preserving the 20-class output head."
    )
    add_bullets(
        doc,
        [
            "Supported classes: cat, dog, horse, cow, sheep, goat, pig, rabbit, chicken, duck, goose, deer, monkey, fox, wolf, bear, tiger, lion, zebra, giraffe.",
            "Data sources include manually provided labels, converted JSON labels, auto-label previews, and targeted fine-tuning datasets for weak classes.",
            "The final dataset path is dataset/animal_dataset_6sp_yolo_20cls, with training and validation splits configured by configs/animal_dataset_6sp_20cls.yaml.",
        ],
    )

    add_paragraph(doc, "Model training and fine-tuning", "Heading 2")
    add_paragraph(
        doc,
        "The base detector is YOLO11s with COCO pretraining. The first custom stage trained animal_dataset1 as a 20-class model. "
        "The model was then fine-tuned on animal_dataset_5sp_v1 and continued on animal_dataset_6sp_v1. During these later stages the model remained a 20-class detector, which avoided losing the original class vocabulary while improving several weak categories."
    )
    add_table(
        doc,
        ["Stage", "Input data", "Purpose", "Output model"],
        [
            ["Base", "YOLO11s COCO weights", "General object localization prior", "yolo11s.pt"],
            ["Stage 1", "animal_dataset1", "Train 20 animal classes", "animal_dataset1 best.pt"],
            ["Stage 2", "animal_dataset_5sp_v1", "Improve wolf, goose, horse, cow, monkey related recognition", "5sp 20-class best.pt"],
            ["Stage 3", "animal_dataset_6sp_v1", "Continue optimization from best model", "6sp 20-class best.pt"],
        ],
        [2.4, 3.5, 5.2, 4.2],
    )
    add_paragraph(
        doc,
        "The final training used CUDA on an NVIDIA GeForce RTX 3070 Ti Laptop GPU, image size 640, batch size 8, freeze=10, early stopping patience=15, and augmentation including HSV jitter, translation, scaling, flipping, mosaic, and light mixup."
    )

    add_paragraph(doc, "Detection and post-processing workflow", "Heading 2")
    add_bullets(
        doc,
        [
            "Run YOLO prediction for each image with configurable confidence threshold, IoU threshold, image size, and CUDA device.",
            "Map predicted labels to the 20 allowed animal categories and discard unsupported labels.",
            "Apply duplicate suppression after YOLO NMS: same-class boxes are merged when overlapping or very close; cross-class boxes are merged only when they likely cover the same object.",
            "Count remaining detections by category and write a JSON dictionary for each image.",
            "For visual inspection, copy validation images and draw bounding boxes with class names and confidence scores.",
        ],
    )

    add_paragraph(doc, "Evaluation Result", "Heading 1")
    add_paragraph(
        doc,
        "The final internal validation used the 15-image validation set discussed during development. The selected post-processing configuration is the balanced suppression version, which reduces duplicate boxes without aggressively deleting nearby animals. The average score under the project rule is 90.94/100, with val_10 counted as full score according to the manual correction."
    )
    add_table(
        doc,
        ["Metric / Result", "Value"],
        [
            ["Final validation score", "90.94 / 100"],
            ["Final model mAP50", "0.978"],
            ["Final model mAP50-95", "0.940"],
            ["Final precision / recall", "0.963 / 0.977"],
            ["Best validation epoch", "Epoch 2 of the 6sp fine-tuning stage"],
            ["Final model path", "runs/detect/runs/detect/animal_dataset_6sp_20cls_finetune/weights/best.pt"],
        ],
        [5.2, 8.8],
    )
    add_table(
        doc,
        ["Image group", "Representative results"],
        [
            ["Perfect predictions", "eval_001384, eval_001388, eval_001431, val_4, val_5, val_6, val_7, val_8, val_10"],
            ["Main misses", "val_1 missed fox; eval_001391 predicted cow count as 2 instead of 3"],
            ["Main confusions", "val_2 and val_9 confused duck with goose; eval_001402 produced extra horse detections"],
        ],
        [3.5, 10.5],
    )
    add_figure(
        doc,
        PROJECT_ROOT / "runs" / "detect" / "runs" / "detect" / "animal_dataset_6sp_20cls_finetune" / "results.png",
        "Figure 1. Training and validation curves from the final 6sp 20-class fine-tuning stage.",
        5.8,
    )
    add_figure(
        doc,
        PROJECT_ROOT / "outputs" / "val_set_annotated_6sp_20cls_finetuned_suppressed_balanced" / "annotated" / "eval_001384.png",
        "Figure 2. Example validation output with class labels, confidence scores, and bounding boxes.",
        4.9,
    )

    add_paragraph(doc, "Analysis", "Heading 1")
    add_paragraph(
        doc,
        "The model performs well on large, visually distinctive animals such as zebra, giraffe, horse, lion, tiger, and bear. The main remaining errors are caused by visually similar species and occlusion. Duck and goose are especially difficult because they share body shape and color patterns in synthetic scenes. Fox, wolf, and dog can also overlap in texture and pose. Count errors usually occur when animals are partially hidden or when two detections are produced for one object."
    )
    add_bullets(
        doc,
        [
            "Duplicate boxes were handled by an extra post-processing layer after YOLO NMS. The balanced setting keeps the highest-confidence box for likely duplicates while avoiding excessive removal of nearby true animals.",
            "Class confusion was reduced by targeted fine-tuning, but duck/goose and dog/wolf/fox remain important future data-collection targets.",
            "False positive categories are costly because the project rule deducts five points per extra category, so threshold tuning must balance recall and category precision.",
        ],
    )

    add_paragraph(doc, "Contributions", "Heading 1")
    add_table(
        doc,
        ["Member", "Contribution"],
        [
            ["待填写", "Dataset conversion, JSON/LabelMe annotation workflow, and validation data organization."],
            ["待填写", "YOLO training, CUDA environment setup, model fine-tuning, and metric tracking."],
            ["待填写", "Evaluation scripts, duplicate-box post-processing, report, PPT, and error analysis."],
        ],
        [3.0, 11.0],
    )

    add_paragraph(doc, "Conclusion", "Heading 1")
    add_paragraph(
        doc,
        "The project completed a reproducible YOLO-based animal detection and counting pipeline covering data preparation, training, fine-tuning, batch inference, JSON output validation, and visual inspection. The final 20-class model achieved strong internal validation performance and produced annotated validation images with probabilities and bounding boxes. Future improvements should focus on adding more manually verified examples for duck/goose, fox/wolf/dog, and partially occluded multi-animal scenes."
    )

    doc.save(REPORT_PATH)
    print(REPORT_PATH)


if __name__ == "__main__":
    main()
