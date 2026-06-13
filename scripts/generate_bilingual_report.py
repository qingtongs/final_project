"""
Generate bilingual (Chinese/English) report based on 报告模板.docx.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = PROJECT_ROOT / "项目要求" / "报告模板.docx"
OUTPUT_DIR = PROJECT_ROOT / "outputs" / "final_deliverables"
REPORT_PATH = OUTPUT_DIR / "animal_detection_project_report.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def remove_from_paragraph(doc: Document, start_index: int) -> None:
    body = doc._body._element
    for paragraph in list(doc.paragraphs[start_index:]):
        body.remove(paragraph._element)


def set_document_defaults(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size, color in [
        ("Heading 1", 15, "1F4E79"),
        ("Heading 2", 12, "2F5597"),
    ]:
        style = doc.styles[style_name] if style_name in [s.name for s in doc.styles] else doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(4)

    if "List Bullet" not in [s.name for s in doc.styles]:
        bullet = doc.styles.add_style("List Bullet", WD_STYLE_TYPE.PARAGRAPH)
        bullet.font.name = "Arial"
        bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        bullet.font.size = Pt(10.5)


def add_paragraph(doc: Document, text: str = "", style: str | None = None,
                  bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.name = "Arial"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    for run in p.runs:
        run.font.name = "Arial"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return p


def add_bilingual_paragraph(doc: Document, en_text: str, zh_text: str,
                            style: str | None = None):
    """Add a bilingual paragraph: English first (bold), then Chinese."""
    p = doc.add_paragraph(style=style)
    # English part
    r_en = p.add_run(en_text)
    r_en.font.name = "Arial"
    r_en.bold = False
    r_en.font.size = Pt(10.5)
    # Separator + Chinese
    r_zh = p.add_run(f"\n{zh_text}")
    r_zh.font.name = "Microsoft YaHei"
    r_zh.font.size = Pt(9.5)
    r_zh.font.color.rgb = RGBColor(0x5E, 0x6A, 0x7D)
    return p


def add_bullets(doc: Document, items: list[tuple[str, str]]) -> None:
    """Add bilingual bullet points. Each item is (en_text, zh_text)."""
    for en_text, zh_text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5)
        r_en = p.add_run(en_text)
        r_en.font.name = "Arial"
        r_en.font.size = Pt(10.5)
        r_zh = p.add_run(f"\n    {zh_text}")
        r_zh.font.name = "Microsoft YaHei"
        r_zh.font.size = Pt(9)
        r_zh.font.color.rgb = RGBColor(0x5E, 0x6A, 0x7D)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(header_cells[i], header, bold=True)
        set_cell_shading(header_cells[i], "D9EAF7")
        header_cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].width = Cm(widths[i])


def add_figure(doc: Document, image_path: Path, caption_en: str, caption_zh: str,
               width: float = 5.8) -> None:
    if not image_path.exists():
        add_paragraph(doc, f"[Image not found: {image_path.name}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    # Bilingual caption
    c = doc.add_paragraph()
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_en = c.add_run(caption_en)
    r_en.italic = True
    r_en.font.size = Pt(9)
    r_en.font.name = "Arial"
    r_zh = c.add_run(f"\n{caption_zh}")
    r_zh.italic = True
    r_zh.font.size = Pt(8)
    r_zh.font.name = "Microsoft YaHei"
    r_zh.font.color.rgb = RGBColor(0x5E, 0x6A, 0x7D)


def configure_cover(doc: Document) -> None:
    """Configure the cover page with bilingual content."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == "Deep Learning Group Project":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = RGBColor(31, 78, 121)
                run.font.name = "Arial"
        elif text.startswith("(2025-2026"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif text.startswith("提交日期"):
            paragraph.text = "Submission Date / 提交日期：2026年 6月 日"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add project title on cover
    # Insert a new paragraph after the course line for the project title
    if doc.paragraphs:
        # Find the date paragraph and insert after it
        for i, p in enumerate(doc.paragraphs):
            if "提交日期" in p.text or "Submission Date" in p.text:
                # Insert project title after the date paragraph
                body = doc._body._element
                new_p = OxmlElement("w:p")
                pPr = OxmlElement("w:pPr")
                jc = OxmlElement("w:jc")
                jc.set(qn("w:val"), "center")
                pPr.append(jc)
                new_p.append(pPr)
                r = OxmlElement("w:r")
                rPr = OxmlElement("w:rPr")
                rFonts = OxmlElement("w:rFonts")
                rFonts.set(qn("w:ascii"), "Arial")
                rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                rPr.append(rFonts)
                sz = OxmlElement("w:sz")
                sz.set(qn("w:val"), "32")  # 16pt
                rPr.append(sz)
                b = OxmlElement("w:b")
                rPr.append(b)
                color = OxmlElement("w:color")
                color.set(qn("w:val"), "1F4E79")
                rPr.append(color)
                r.append(rPr)
                t = OxmlElement("w:t")
                t.text = "Animal Detection and Counting / 动物检测与计数"
                t.set(qn("xml:space"), "preserve")
                r.append(t)
                new_p.append(r)
                p._element.addnext(new_p)
                break

    # Configure info table
    if doc.tables:
        table = doc.tables[0]
        values = [
            ["姓   名\nName", "待填写\nTBD", "学    号\nStudent ID", "待填写\nTBD"],
            ["姓   名\nName", "待填写\nTBD", "学    号\nStudent ID", "待填写\nTBD"],
            ["姓   名\nName", "待填写\nTBD", "学    号\nStudent ID", "待填写\nTBD"],
            ["专业班级\nClass", "待填写\nTBD", "课程代码\nCourse Code", "待填写\nTBD"],
            ["课程名称\nCourse", "Deep Learning", "任课教师\nInstructor", "曾子倩"],
        ]
        for row, values_row in zip(table.rows, values):
            for cell, value in zip(row.cells, values_row):
                set_cell_text(cell, value,
                             bold=value.split("\n")[0] in {
                                 "姓   名", "学    号", "专业班级", "课程代码", "课程名称", "任课教师",
                                 "Name", "Student ID", "Class", "Course Code", "Course", "Instructor"
                             },
                             font_size=9)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document(str(TEMPLATE_PATH))
    set_document_defaults(doc)
    configure_cover(doc)

    # Find the "Introduction" paragraph to start replacing content
    try:
        start = next(i for i, p in enumerate(doc.paragraphs)
                     if p.text.strip().startswith("Introduction"))
    except StopIteration:
        start = len(doc.paragraphs)
    remove_from_paragraph(doc, start)

    # ============================
    # 1. INTRODUCTION / 引言
    # ============================
    doc.add_page_break()
    add_paragraph(doc, "1. Introduction / 引言", "Heading 1")

    add_bilingual_paragraph(
        doc,
        "This project implements an animal detection and counting system for the Deep Learning "
        "final project. The system accepts a single image as input and returns a clean Python-style "
        "dictionary whose keys are supported animal categories and whose values are positive integer "
        "counts. The final implementation follows the required 20-category vocabulary and writes "
        "batch predictions to JSON without any natural-language explanations.",
        "本项目为深度学习期末项目实现了动物检测与计数系统。系统以单张图片作为输入，返回一个干净的Python风格"
        "字典，键为支持的动物类别，值为正整数计数。最终实现遵循要求的20类词汇表，并将批量预测结果写入JSON文件，"
        "不含任何自然语言解释。"
    )
    add_bilingual_paragraph(
        doc,
        "The work is based on YOLO11s and the Ultralytics training/inference framework. YOLO was "
        "selected because it performs object localization, classification, and confidence estimation "
        "in a single forward pass, which is highly suitable for the time-limited validation and final "
        "evaluation workflow. The project covers the complete pipeline: data collection and annotation, "
        "model training and fine-tuning, detection inference, post-processing, and JSON prediction export.",
        "本项目基于YOLO11s和Ultralytics训练/推理框架。选择YOLO是因为它在单次前向传播中同时完成目标定位、"
        "分类和置信度估计，非常适合有时间限制的验证和最终评估流程。项目覆盖了完整流程：数据收集与标注、"
        "模型训练与微调、检测推理、后处理和JSON预测导出。"
    )

    # ============================
    # 2. METHODOLOGY / 方法
    # ============================
    add_paragraph(doc, "2. Methodology / 方法", "Heading 1")

    # 2.1 Dataset Construction
    add_paragraph(doc, "2.1 Dataset Construction / 数据集构建", "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The training data was built progressively through multiple stages. Initially, two source "
        "datasets (animal_dataset_xfz and animal_dataset_yzy) were merged and their annotations were "
        "converted into a unified JSON format. From this merged dataset, animal_dataset1 was prepared "
        "as a 20-class YOLO-format dataset with training and validation splits.",
        "训练数据通过多个阶段逐步构建。首先，合并了两个源数据集（animal_dataset_xfz和animal_dataset_yzy），"
        "并将其标注转换为统一的JSON格式。基于合并后的数据集，准备了animal_dataset1作为20类YOLO格式数据集，"
        "包含训练集和验证集划分。"
    )
    add_bilingual_paragraph(
        doc,
        "To improve performance on weak categories, two additional targeted datasets were constructed: "
        "animal_dataset_5sp_v1 (focusing on wolf, goose, horse, cow, and monkey) and "
        "animal_dataset_6sp_v1 (continuing optimization of difficult categories). Both datasets were "
        "auto-labeled in LabelMe-style JSON format and then converted to YOLO format while preserving "
        "the full 20-class output head. This approach ensured that fine-tuning on targeted categories "
        "did not reduce the model's ability to recognize the original 20 animal classes.",
        "为改善弱类别表现，构建了两个额外的定向数据集：animal_dataset_5sp_v1（聚焦狼、鹅、马、牛和猴子）"
        "和animal_dataset_6sp_v1（继续优化困难类别）。两个数据集均以LabelMe风格JSON格式自动标注，然后转换为"
        "YOLO格式，同时保留完整的20类输出头。该方法确保在定向类别上微调不会降低模型识别原始20个动物类别的能力。"
    )
    add_bullets(doc, [
        ("Supported classes: cat, dog, horse, cow, sheep, goat, pig, rabbit, chicken, duck, goose, "
         "deer, monkey, fox, wolf, bear, tiger, lion, zebra, giraffe (20 total).",
         "支持类别：猫、狗、马、牛、羊、山羊、猪、兔、鸡、鸭、鹅、鹿、猴、狐狸、狼、熊、老虎、狮子、斑马、"
         "长颈鹿（共20类）。"),
        ("Data sources include manually provided labels, auto-labeled JSON annotations with visual "
         "preview, and targeted fine-tuning datasets for weak classes.",
         "数据来源包括人工标注、带可视化预览的自动标注JSON，以及针对弱类别的定向微调数据集。"),
        ("Final dataset path: dataset/animal_dataset_6sp_yolo_20cls, configured by "
         "configs/animal_dataset_6sp_20cls.yaml.",
         "最终数据集路径：dataset/animal_dataset_6sp_yolo_20cls，由configs/animal_dataset_6sp_20cls.yaml配置。"),
    ])

    # 2.2 Model Training
    add_paragraph(doc, "2.2 Model Training and Fine-tuning / 模型训练与微调", "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The base detector is YOLO11s (the 'small' variant of YOLO11) with COCO pre-training weights. "
        "Training proceeded in three stages, each building on the previous best model:",
        "基础检测器为YOLO11s（YOLO11的小型变体），使用COCO预训练权重。训练分三个阶段进行，每个阶段基于"
        "前一阶段的最佳模型："
    )
    add_table(doc,
              ["Stage / 阶段", "Input Data / 输入数据", "Purpose / 目的",
               "Output Model / 输出模型"],
              [
                  ["Stage 1", "animal_dataset1", "Train 20 animal classes from COCO baseline\n从COCO基线训练20个动物类别",
                   "animal_dataset1 best.pt"],
                  ["Stage 2", "animal_dataset_5sp_v1", "Improve wolf/goose/horse/cow/monkey\n改善狼/鹅/马/牛/猴的识别",
                   "5sp 20-class best.pt"],
                  ["Stage 3", "animal_dataset_6sp_v1", "Continue optimization from best model\n从最佳模型继续优化",
                   "6sp 20-class best.pt"],
              ],
              [2.0, 3.2, 5.5, 3.5],
              )
    add_bilingual_paragraph(
        doc,
        "Training hyperparameters: image size 640×640, batch size 8, freeze first 10 layers, "
        "early stopping patience 15 epochs, AdamW optimizer with initial learning rate 0.0005, "
        "weight decay 0.0005, 3 warmup epochs. Data augmentations included HSV jitter "
        "(h=0.015, s=0.7, v=0.4), rotation (±8°), translation (±10%), scaling (±40%), "
        "horizontal flip (50%), vertical flip (5%), mosaic (100%), and light mixup (5%). "
        "Mosaic augmentation was disabled after epoch 10. Training was performed on a single "
        "NVIDIA GeForce RTX 3070 Ti Laptop GPU using CUDA.",
        "训练超参数：图像尺寸640×640，批次大小8，冻结前10层，早停耐心值15轮，AdamW优化器（初始学习率"
        "0.0005，权重衰减0.0005，预热3轮）。数据增强包括HSV抖动（h=0.015, s=0.7, v=0.4）、旋转（±8°）、"
        "平移（±10%）、缩放（±40%）、水平翻转（50%）、垂直翻转（5%）、马赛克（100%）和轻量mixup（5%）。"
        "马赛克增强在第10轮后关闭。训练在单张NVIDIA GeForce RTX 3070 Ti笔记本GPU上使用CUDA进行。"
    )

    # 2.3 Detection and Post-processing
    add_paragraph(doc, "2.3 Detection and Post-processing Workflow / 检测与后处理流程",
                  "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The inference pipeline processes each image through the following steps:",
        "推理流程按以下步骤处理每张图片："
    )
    add_bullets(doc, [
        ("1) YOLO Prediction: Run the fine-tuned YOLO11s model to obtain bounding boxes, "
         "confidence scores, and predicted class labels for each detection.",
         "1) YOLO预测：运行微调后的YOLO11s模型，获取每个检测的边界框、置信度分数和预测类别标签。"),
        ("2) Category Whitelist Filtering: Map YOLO output labels to the 20 supported animal "
         "categories. Unsupported labels (e.g., person, bird, elephant) are discarded.",
         "2) 类别白名单过滤：将YOLO输出标签映射到20个支持的动物类别。不支持的标签（如人、鸟、大象）被丢弃。"),
        ("3) Confidence Thresholding: Discard detections with confidence scores below 0.25. "
         "YOLO's built-in NMS is applied with IoU threshold 0.45.",
         "3) 置信度阈值：丢弃置信度低于0.25的检测。YOLO内置NMS以IoU阈值0.45应用。"),
        ("4) Custom Duplicate Suppression: An additional post-processing layer beyond YOLO NMS. "
         "Same-class boxes are merged when IoU ≥ 0.40 or when centers are close (≤30% of min diagonal) "
         "with similar sizes (area ratio ≥ 0.50). Cross-class boxes are merged under stricter "
         "conditions: IoU ≥ 0.62 or centers ≤18% of min diagonal with area ratio ≥ 0.70. "
         "In all cases, the highest-confidence detection is retained.",
         "4) 自定义重复抑制：YOLO NMS之外的额外后处理层。同类别框在IoU≥0.40时合并，或中心距离近"
         "（≤最小对角线的30%）且尺寸相似（面积比≥0.50）时合并。跨类别框在更严格条件下合并："
         "IoU≥0.62或中心距离≤最小对角线的18%且面积比≥0.70。所有情况下保留置信度最高的检测。"),
        ("5) Counting: Group remaining detections by mapped category name and count instances "
         "per category.",
         "5) 计数：按映射后的类别名称分组剩余检测，统计每个类别的实例数。"),
        ("6) JSON Export: Write a dictionary for each image containing only categories with "
         "positive counts. Output format: {\"eval_000001\": {\"cat\": 2, \"duck\": 1, \"deer\": 1}}.",
         "6) JSON导出：为每张图片写入一个字典，仅包含具有正计数的类别。输出格式："
         "{\"eval_000001\": {\"cat\": 2, \"duck\": 1, \"deer\": 1}}。"),
    ])
    add_bilingual_paragraph(
        doc,
        "The same duplicate-suppression logic is shared between JSON prediction counting and "
        "annotated image visualization, ensuring consistency between displayed boxes and submitted counts.",
        "相同的重复抑制逻辑在JSON预测计数和标注图像可视化之间共享，确保显示框与提交计数之间的一致性。"
    )

    # ============================
    # 3. EVALUATION RESULT / 评估结果
    # ============================
    add_paragraph(doc, "3. Evaluation Result / 评估结果", "Heading 1")

    # Final model metrics
    add_paragraph(doc, "3.1 Final Model Performance / 最终模型性能", "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The final model achieved strong performance across all training metrics after the third "
        "fine-tuning stage. Epoch 2 was selected as the best checkpoint based on validation mAP50.",
        "最终模型在第三阶段微调后，所有训练指标均表现强劲。基于验证集mAP50，第2轮被选为最佳检查点。"
    )
    add_table(doc,
              ["Metric / 指标", "Value / 数值"],
              [
                  ["Precision / 精确率", "0.963"],
                  ["Recall / 召回率", "0.977"],
                  ["mAP50", "0.978"],
                  ["mAP50-95", "0.940"],
                  ["Best epoch / 最佳轮次", "Epoch 2"],
                  ["Model path / 模型路径",
                   "runs/detect/.../animal_dataset_6sp_20cls_finetune/weights/best.pt"],
              ],
              [4.5, 9.5],
              )

    # Validation results
    add_paragraph(doc, "3.2 Validation Set Results / 验证集结果", "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The final internal validation used the 15-image validation set (val_set). Using the "
        "balanced suppression post-processing configuration, the model achieved an average score "
        "of 90.94/100 under the project scoring rules. 9 out of 15 images received perfect scores "
        "(100/100), demonstrating that the system correctly handles most detection scenarios.",
        "最终内部验证使用15张图片的验证集（val_set）。使用平衡式抑制后处理配置，模型在项目评分规则下取得"
        "了平均90.94/100的评分。15张图片中有9张获得满分（100/100），表明系统正确处理了大多数检测场景。"
    )

    add_table(doc,
              ["Image Group / 图片组", "Representative Results / 代表性结果"],
              [
                  ["Perfect predictions / 完美预测",
                   "eval_001384, eval_001388, eval_001431, val_4, val_5, val_6, val_7, val_8, val_10"],
                  ["Main misses / 主要漏检",
                   "val_1 missed fox / val_1漏检狐狸; eval_001391 predicted cow count as 2 instead of 3 / "
                   "eval_001391将牛的数量预测为2而非3"],
                  ["Main confusions / 主要混淆",
                   "val_2 and val_9 confused duck with goose / val_2和val_9将鸭与鹅混淆; "
                   "eval_001402 produced extra horse detections / eval_001402产生多余的马检测"],
              ],
              [3.5, 10.5],
              )

    # Training curves figure
    add_figure(
        doc,
        PROJECT_ROOT / "runs" / "detect" / "runs" / "detect" / "animal_dataset_6sp_20cls_finetune" / "results.png",
        "Figure 1. Training and validation curves from the final 6sp 20-class fine-tuning stage.",
        "图1. 最终6sp 20类微调阶段的训练和验证曲线。",
        5.5,
    )

    # Annotated example
    add_figure(
        doc,
        PROJECT_ROOT / "outputs" / "val_set_annotated_6sp_20cls_finetuned_suppressed_balanced" / "annotated" / "eval_001384.png",
        "Figure 2. Example validation output with class labels, confidence scores, and bounding boxes.",
        "图2. 带类别标签、置信度分数和边界框的验证输出示例。",
        4.5,
    )

    # ============================
    # 4. ANALYSIS / 分析
    # ============================
    add_paragraph(doc, "4. Analysis / 分析", "Heading 1")

    add_paragraph(doc, "4.1 Performance Analysis / 性能分析", "Heading 2")
    add_bilingual_paragraph(
        doc,
        "The model performs excellently on large, visually distinctive animals such as zebra, "
        "giraffe, horse, lion, tiger, and bear, where mAP exceeds 0.95. These categories benefit "
        "from distinctive shape, texture, and color patterns that YOLO learns well from the "
        "training data.",
        "模型在大型、视觉特征明显的动物（如斑马、长颈鹿、马、狮子、老虎和熊）上表现优异，mAP超过0.95。"
        "这些类别受益于独特的形状、纹理和颜色模式，YOLO能够从训练数据中很好地学习。"
    )
    add_bilingual_paragraph(
        doc,
        "The main remaining errors stem from three sources: (1) visually similar species, "
        "particularly duck vs. goose which share body shape and color in synthetic scenes, and "
        "fox vs. wolf vs. dog which overlap in texture and pose; (2) occlusion and tightly grouped "
        "animals where partial visibility leads to missed or merged detections; and (3) synthetic "
        "scene artifacts that may not perfectly represent real-world animal appearances.",
        "剩余错误主要来自三个方面：（1）视觉相似物种，特别是鸭和鹅在合成场景中体型和颜色相似，以及狐狸、"
        "狼和狗在纹理和姿态上重叠；（2）遮挡和密集排列动物，部分可见性导致漏检或合并检测；"
        "（3）合成场景伪影可能无法完美代表真实世界的动物外观。"
    )

    add_paragraph(doc, "4.2 Error Analysis / 错误分析", "Heading 2")
    add_bullets(doc, [
        ("Duck/Goose Confusion: These two bird species are the most commonly confused pair. "
         "In synthetic images, both appear with similar white/brown body coloration and comparable "
         "size, making fine-grained discrimination difficult without more targeted training examples.",
         "鸭/鹅混淆：这两种禽类是最常混淆的组合。在合成图像中，两者呈现相似的白色/棕色体色和相近大小，"
         "没有更多定向训练样本的情况下，细粒度区分非常困难。"),
        ("Canine Confusion: Fox, wolf, and dog share similar body plans, fur textures, and poses. "
         "When depicted at a distance or in non-distinctive poses, the model may incorrectly classify "
         "among these three categories.",
         "犬科混淆：狐狸、狼和狗具有相似的身体结构、皮毛纹理和姿态。在远距离或非特征性姿态下，"
         "模型可能在这三个类别之间错误分类。"),
        ("Counting Errors: Occlusion and tight grouping remain challenging. When animals partially "
         "overlap, the model may detect two animals as one, or one animal as two if body parts are "
         "visually separated by occluders.",
         "计数错误：遮挡和密集排列仍然具有挑战性。当动物部分重叠时，模型可能将两只动物检测为一只，"
         "或当身体部位被遮挡物视觉分离时将一只动物检测为两只。"),
        ("False Positive Penalties: The scoring rule deducts 5 points per extra wrong category, "
         "making false positive categories especially costly. Conservative confidence thresholding "
         "is needed to balance recall against the risk of introducing hallucinated categories.",
         "误报惩罚：评分规则对每个额外错误类别扣除5分，使误报类别代价特别高。需要保守的置信度阈值"
         "来平衡召回率和引入虚假类别的风险。"),
    ])

    add_paragraph(doc, "4.3 Improvement Strategies / 改进策略", "Heading 2")
    add_bullets(doc, [
        ("Add more manually verified training examples specifically for duck/goose and "
         "fox/wolf/dog pairs, with careful annotation review to ensure labeling consistency.",
         "专门为鸭/鹅和狐狸/狼/狗组合增加更多人工验证的训练样本，并进行仔细的标注审查以确保标注一致性。"),
        ("Explore test-time augmentation (TTA) during inference to improve robustness, "
         "particularly for images with occlusion or unusual viewpoints.",
         "在推理过程中探索测试时增强（TTA）以提高鲁棒性，特别是针对有遮挡或不寻常视角的图像。"),
        ("Investigate per-class confidence threshold tuning: higher thresholds for frequently "
         "confused categories and lower thresholds for reliably detected categories.",
         "研究逐类置信度阈值调优：对经常混淆的类别使用更高阈值，对可靠检测的类别使用更低阈值。"),
        ("Consider ensemble approaches combining multiple YOLO variants or incorporating "
         "additional feature extractors for fine-grained species discrimination.",
         "考虑集成方法，结合多个YOLO变体或引入额外的特征提取器进行细粒度物种区分。"),
    ])

    # ============================
    # 5. CONTRIBUTIONS / 贡献
    # ============================
    add_paragraph(doc, "5. Contributions / 贡献", "Heading 1")
    add_bilingual_paragraph(
        doc,
        "The following table describes each team member's contribution to the project.",
        "下表描述了每位团队成员对项目的贡献。"
    )
    add_table(doc,
              ["Member / 成员", "Contributions / 贡献"],
              [
                  ["Member 1 / 成员1\n待填写 / TBD",
                   "Dataset collection, merging, and conversion to LabelMe JSON annotation format. "
                   "Validation data organization and preprocessing scripts.\n"
                   "数据集收集、合并和LabelMe JSON标注格式转换。验证数据组织和预处理脚本。"],
                  ["Member 2 / 成员2\n待填写 / TBD",
                   "YOLO model training pipeline, CUDA environment setup, multi-stage fine-tuning, "
                   "hyperparameter tuning, and model performance tracking.\n"
                   "YOLO模型训练流程、CUDA环境配置、多阶段微调、超参数调优和模型性能跟踪。"],
                  ["Member 3 / 成员3\n待填写 / TBD",
                   "Evaluation and inference scripts, custom duplicate-box post-processing algorithm, "
                   "annotated image generation, report writing, PPT creation, and error analysis.\n"
                   "评估和推理脚本、自定义重复框后处理算法、标注图像生成、报告撰写、PPT制作和错误分析。"],
              ],
              [3.0, 11.0],
              )

    # ============================
    # 6. CONCLUSION / 结论
    # ============================
    add_paragraph(doc, "6. Conclusion / 结论", "Heading 1")
    add_bilingual_paragraph(
        doc,
        "This project successfully completed a reproducible YOLO-based animal detection and "
        "counting pipeline covering the full workflow: data preparation and annotation, multi-stage "
        "model fine-tuning, batch inference with custom post-processing, JSON prediction export, "
        "validation scoring, and visual annotation with bounding boxes and confidence scores. "
        "The final 20-class model achieved a high validation score of 90.94/100 and strong "
        "detection metrics (mAP50: 0.978, mAP50-95: 0.940).",
        "本项目成功完成了一个可复现的基于YOLO的动物检测与计数流程，覆盖了完整工作流：数据准备与标注、"
        "多阶段模型微调、带自定义后处理的批量推理、JSON预测导出、验证评分以及带边界框和置信度分数的"
        "可视化标注。最终20类模型取得了90.94/100的高验证评分和强劲的检测指标（mAP50: 0.978, mAP50-95: 0.940）。"
    )
    add_bilingual_paragraph(
        doc,
        "The key technical contributions include: (1) a progressive data construction strategy that "
        "targets weak categories while preserving the full 20-class vocabulary; (2) a balanced "
        "duplicate suppression algorithm that reduces false positives without aggressive deletion "
        "of true nearby animals; and (3) shared post-processing logic between visual annotation "
        "and JSON counting to ensure consistency between displayed and submitted results.",
        "关键技术贡献包括：（1）渐进式数据构建策略，针对弱类别同时保留完整20类词汇表；"
        "（2）平衡式重复抑制算法，减少误报但不过度删除真实邻近动物；（3）可视化标注和JSON计数之间"
        "共享后处理逻辑，确保显示结果与提交结果一致。"
    )
    add_bilingual_paragraph(
        doc,
        "Future work should focus on collecting more manually verified hard examples for the most "
        "confusing category pairs (duck/goose, fox/wolf/dog), exploring per-class threshold "
        "optimization, and investigating advanced techniques such as test-time augmentation and "
        "ensemble methods to further improve robustness on challenging cases involving occlusion "
        "and fine-grained species discrimination.",
        "未来工作应重点关注为最易混淆的类别对（鸭/鹅、狐狸/狼/狗）收集更多人工验证的困难样本，"
        "探索逐类阈值优化，并研究测试时增强和集成方法等先进技术，以进一步提高在涉及遮挡和细粒度"
        "物种区分的挑战性案例上的鲁棒性。"
    )

    # ============================
    # SAVE
    # ============================
    doc.save(str(REPORT_PATH))
    print(f"Report saved to: {REPORT_PATH}")


if __name__ == "__main__":
    main()
